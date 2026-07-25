#!/usr/bin/env python3
"""Build the authoritative Microsoft MarkItDown MCP security assessment."""

from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parent
STAGES = ROOT / "stages"
OUT = ROOT.parent / "Microsoft_MarkItDown_MCP_VSCode_VisualStudio_v0.0.1a4_Assessment_2026-07-25_v1.0.docx"
FIG = ROOT / "evidence" / "markitdown-trust-boundaries.png"
BLUE, TEAL, LIGHT, PALE, RED = "17365D", "0F6B78", "D9EAF0", "EEF4F7", "C00000"

REFERENCES = [
    ("REF-001", "Microsoft MarkItDown repository", "https://github.com/microsoft/markitdown"),
    ("REF-002", "PyPI: markitdown-mcp 0.0.1a4", "https://pypi.org/project/markitdown-mcp/0.0.1a4/"),
    ("REF-003", "PyPI: markitdown 0.1.6", "https://pypi.org/project/markitdown/0.1.6/"),
    ("REF-004", "VS Code MCP servers documentation", "https://code.visualstudio.com/docs/agent-customization/mcp-servers"),
    ("REF-005", "Visual Studio MCP servers documentation", "https://learn.microsoft.com/en-us/visualstudio/ide/mcp-servers?view=visualstudio"),
    ("REF-006", "MCP Python SDK DNS rebinding advisory", "https://github.com/modelcontextprotocol/python-sdk/security/advisories/GHSA-9h52-p55h-vw2f"),
    ("REF-007", "MCP Python SDK malformed-request DoS advisory", "https://github.com/modelcontextprotocol/python-sdk/security/advisories/GHSA-3qhf-m339-9g5v"),
    ("REF-008", "MCP Python SDK Streamable HTTP DoS advisory", "https://github.com/modelcontextprotocol/python-sdk/security/advisories/GHSA-j975-95f5-7wqh"),
    ("REF-009", "OWASP Top 10 for LLM Applications 2025", "https://owasp.org/www-project-top-10-for-large-language-model-applications/assets/PDF/OWASP-Top-10-for-LLMs-v2025.pdf"),
    ("REF-010", "OWASP Agentic AI Threats and Mitigations", "https://genai.owasp.org/resource/agentic-ai-threats-and-mitigations/"),
    ("REF-011", "OWASP ASVS 5.0", "https://github.com/OWASP/ASVS/releases/tag/v5.0.0_release"),
    ("REF-012", "MITRE ATLAS", "https://atlas.mitre.org/"),
    ("REF-013", "NIST AI Risk Management Framework", "https://www.nist.gov/itl/ai-risk-management-framework"),
    ("REF-014", "NIST AI 600-1 Generative AI Profile", "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf"),
    ("REF-015", "NCSC Guidelines for Secure AI System Development", "https://www.ncsc.gov.uk/collection/guidelines-secure-ai-system-development/guidelines/secure-development"),
    ("REF-016", "CIS Critical Security Controls v8.1", "https://www.cisecurity.org/controls"),
    ("REF-017", "ICO guidance on AI and data protection", "https://ico.org.uk/for-organisations/uk-gdpr-guidance-and-resources/artificial-intelligence/"),
    ("REF-018", "FCA operational resilience", "https://www.fca.org.uk/firms/operational-resilience"),
    ("REF-019", "PRA SS2/21 outsourcing and third-party risk", "https://www.bankofengland.co.uk/prudential-regulation/publication/2021/march/outsourcing-and-third-party-risk-management-ss"),
    ("REF-020", "UK Data Protection Act 2018", "https://www.legislation.gov.uk/ukpga/2018/12/contents"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:tblHeader"))
    properties.append(OxmlElement("w:cantSplit"))


def prevent_row_split(row):
    row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = value
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
    repeat_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row):
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if ridx % 2:
                set_cell_shading(cells[i], PALE)
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
    doc.add_paragraph()
    return table


def add_hyperlink(paragraph, text, url=None, anchor=None):
    hyperlink = OxmlElement("w:hyperlink")
    if url:
        rel_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink.set(qn("r:id"), rel_id)
    if anchor:
        hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([rpr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_field(paragraph, instruction):
    run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run.extend([begin, instr, separate])
    result = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    result.append(text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    paragraph._p.extend([run, result, end_run])


def cite(paragraph, number):
    paragraph.add_run(" ")
    add_hyperlink(paragraph, f"[REF-{number:03d}]", anchor=f"REF_{number:03d}")


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_narrative(doc, title, paragraphs):
    doc.add_heading(title, level=2)
    for item in paragraphs:
        doc.add_paragraph(item)


def make_figure():
    image = Image.new("RGB", (1500, 700), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default(size=28)
    small = ImageFont.load_default(size=20)
    boxes = [
        (40, 250, 280, 450, "Untrusted input\nfile / URL / data URI", LIGHT),
        (360, 90, 660, 270, "IDE + agent\nVS Code / Visual Studio", "DDEBF7"),
        (360, 420, 660, 600, "MCP transport\nSTDIO or HTTP/SSE", "FFF2CC"),
        (750, 250, 1040, 450, "markitdown-mcp\nconvert_to_markdown(uri)", "FCE4D6"),
        (1140, 120, 1450, 300, "Local filesystem\nand archives", "F4CCCC"),
        (1140, 400, 1450, 580, "External/internal network\nand returned content", "F4CCCC"),
    ]
    for x1, y1, x2, y2, label, fill in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=f"#{fill}", outline=f"#{BLUE}", width=4)
        draw.multiline_text((x1 + 18, y1 + 45), label, fill="black", font=small, spacing=8)
    for start, end in [((280,350),(360,180)),((280,350),(360,510)),((660,180),(750,330)),((660,510),(750,370)),((1040,330),(1140,210)),((1040,370),(1140,490))]:
        draw.line((start, end), fill=f"#{TEAL}", width=7)
    draw.text((40, 30), "MarkItDown MCP trust boundaries and authority expansion", fill=f"#{BLUE}", font=font)
    image.save(FIG)


def configure(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05
    for name, size, color in [("Title", 28, BLUE), ("Heading 1", 18, BLUE), ("Heading 2", 13, TEAL), ("Heading 3", 10.5, BLUE)]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    section = doc.sections[0]
    section.top_margin, section.bottom_margin = Inches(0.85), Inches(0.65)
    section.left_margin, section.right_margin = Inches(0.72), Inches(0.72)
    header = section.header.paragraphs[0]
    header.text = "INTERNAL-CONFIDENTIAL  |  Microsoft MarkItDown MCP Security Assessment"
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SoloSentry Assessment Environment  |  25 July 2026  |  Page ")
    add_field(footer, " PAGE ")
    props = doc.core_properties
    props.author = props.last_modified_by = "SoloSentry Assessment Environment"
    props.title = "Microsoft MarkItDown MCP Security Assessment"
    props.subject = "VS Code and Visual Studio assessment for UK financial services"
    props.keywords = "MarkItDown, MCP, VS Code, Visual Studio, security assessment"


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("MICROSOFT MARKITDOWN MCP")
    r.bold, r.font.size = True, Pt(17)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Security Assessment")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("VS Code on Windows 11 and macOS\nVisual Studio on Windows 11").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nDecision: DO NOT APPROVE FOR PRODUCTION")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nClassification: INTERNAL-CONFIDENTIAL\nVersion 1.0  |  25 July 2026\nUK financial-services use case")
    doc.add_page_break()


def ide_part(doc, title, ide, platforms, ref_num):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(f"{ide} supports configuring and invoking MCP servers")
    cite(p, ref_num)
    doc.add_heading(f"{ide} Decision and Approval", level=2)
    doc.add_paragraph("Decision: Do not approve for production. A future time-limited, non-production STDIO-only pilot may be reconsidered only after the critical URI authority and dependency findings are remediated and host verification succeeds. HTTP and SSE modes remain prohibited.")
    add_narrative(doc, f"{ide} Purpose and Function", [f"The integration makes MarkItDown document conversion available to an AI agent in {ide}. The business purpose is content normalization, but the server’s broad URI interface means a prompt, repository instruction, document or model decision can select sensitive local or network content."])
    add_narrative(doc, f"{ide} Architecture and Trust Boundaries", [
        f"The trust chain is user and workspace → {ide} agent → MCP transport → Python server → converters, filesystem and network → Markdown response → model context. Boundaries include workspace trust, child-process execution, local file access, proxy and DNS routing, archive parsing and any upstream model data boundary.",
        f"On {platforms}, the server inherits the launching user’s permissions unless operating-system, IDE or endpoint controls reduce them. IDE consent is useful but is not a substitute for server-side URI authorization.",
    ])
    add_narrative(doc, f"{ide} Data Flows and Privacy", [
        "Input includes the URI and the bytes reachable through it. Output may contain personal data, customer information, credentials embedded in documents, trading or operational data, source code and document metadata. Converted content is sent back to the IDE and can enter model context, logs, chat history or diagnostics.",
        "No product-specific privacy notice or telemetry control was established for the MCP package. Static code inspection did not identify explicit package telemetry; runtime verification was not observed. Enterprise policy must independently control IDE and model retention and cross-border processing.",
    ])
    doc.add_heading(f"{ide} Installation Manifest and Uninstall", level=2)
    add_table(doc, ["Element", "Required state", "Uninstall verification"], [
        ("Package source", "Approved internal Python repository; hashes pinned", "Package absent and cache disposition recorded"),
        ("MCP configuration", "User/workspace scope documented; STDIO only", "Configuration removed from all scopes"),
        ("Python environment", "Dedicated least-privilege virtual environment", "Environment removed without affecting unrelated packages"),
        ("Logs and converted content", "Approved storage and retention", "Residual content located and disposed under policy"),
    ])
    add_narrative(doc, f"{ide} Runtime, Activation and Network", [
        "Activation occurs when the IDE starts or invokes the configured MCP server. STDIO limits the transport to a child-process channel but does not restrict the tool’s outbound HTTP requests or local file reads. Published HTTP and SSE modes introduce a listening service, are unauthenticated and must not be enabled.",
        "Network egress should be denied by default and brokered through an authenticated proxy with public-host allowlists, private, link-local, loopback and metadata address denial, DNS rebinding defense, redirect revalidation, TLS policy and request timeouts. Verification must cover direct, proxy, IPv4, IPv6 and redirect paths.",
    ])
    add_narrative(doc, f"{ide} Authentication and Authorization", [
        "The tool exposes no caller authentication or per-resource authorization. In STDIO mode, process ancestry provides a limited channel boundary, while the actual file and network authority remains the Python process identity. In HTTP or SSE mode the lack of authentication combines with vulnerable SDK behavior and is unacceptable.",
        "Authorization must be enforced in a remediated wrapper: approved roots, resolved-path containment, file-type controls, public destination policy, maximum redirects, denied schemes, request budgets and explicit human confirmation for sources outside the active workspace.",
    ])
    add_narrative(doc, f"{ide} Telemetry and Logging", [
        "Static analysis found no explicit MarkItDown MCP telemetry client. This is not proof that the full IDE, Python runtime, dependencies or model service produce no diagnostics. Runtime network capture, process monitoring and log review are required.",
        "Security logging should record the normalized source class, policy decision, destination category, converter, byte count, duration, outcome and correlation identifier without recording confidential content, full query strings or credentials.",
    ])
    add_narrative(doc, f"{ide} MCP and Agent Skills", [
        "The integration is an MCP tool, not an installed Agent Skill. Repository instructions and installed skills can nevertheless influence when an agent calls the tool and which URI it selects. Treat all workspace instructions, prompts and documents as untrusted control-plane input.",
        "Controls should require tool allowlisting, visible invocation, user confirmation for sensitive actions, instruction-source provenance, workspace trust and denial of automatic calls from untrusted repositories. No bundled Agent Skill was found in the published wheel.",
    ])
    add_narrative(doc, f"{ide} Supply Chain and Dependencies", [
        "The pre-release MCP package depends on mcp~=1.8.0, resolving to 1.8.1 in the assessment lock. That constraint prevents automatic adoption of the relevant fixes at 1.9.4, 1.10.0 and 1.23.0. The dependency graph is broad because document conversion enables multiple parsers.",
        "An enterprise build must use immutable artifact hashes, a controlled index, complete SBOM, advisory gates, source attestation, reproducible build evidence, quarantine before promotion and a tested update and rollback process.",
    ])
    add_narrative(doc, f"{ide} Threats, Findings and Residual Risk", [
        "F-001 local-file disclosure and SSRF, F-002 unauthenticated vulnerable HTTP transports, F-003 resource exhaustion, F-004 dependency exposure, F-005 missing representative runtime evidence, F-006 regulated-data disclosure, and F-007 MCP source-to-package provenance gap all apply.",
        "Residual risk remains High or Critical under the published configuration. IDE-native controls can reduce accidental activation and process reach, but they cannot correct vulnerable dependencies or enforce URI policy inside this server.",
    ])
    doc.add_heading(f"{ide} Framework Disposition", level=2)
    add_table(doc, ["Framework", "Disposition", "Key mapping"], [
        ("OWASP LLM Top 10", "Applicable", "Prompt injection, improper output handling, excessive agency and supply chain"),
        ("OWASP Agentic / MCP guidance", "Applicable", "Tool authorization, identity, trust boundaries and resource controls"),
        ("OWASP ASVS 5.0", "Applicable by analogy", "Input validation, access control, communications and files"),
        ("MITRE ATLAS", "Applicable", "Agent/tool abuse and exfiltration-oriented scenarios"),
        ("NIST AI RMF / 600-1", "Applicable", "Govern, Map, Measure and Manage evidence"),
        ("NCSC secure AI", "Applicable", "Secure design, development, deployment and maintenance"),
    ])
    doc.add_heading(f"{ide} Controls and Detection", level=2)
    add_bullets(doc, [
        "Prevent: approved internal artifact, exact hashes, STDIO-only configuration, dedicated identity, workspace trust, URI allowlist, egress deny-by-default and resource quotas.",
        "Detect: endpoint process ancestry, unexpected listening sockets, Python child processes, private-address requests, excessive conversions, archive expansion failures and policy denials.",
        "Respond: disable the MCP configuration, terminate the process, isolate the host if disclosure is suspected, preserve sanitized evidence, rotate exposed credentials and follow incident-management policy.",
    ])
    doc.add_heading(f"{ide} Confidence, Limitations, Evidence and Verification", level=2)
    doc.add_paragraph("Confidence is High for package identity, hashes, archive content, source inspection, static analysis, malware results and direct dependency advisories. Confidence is Medium for architecture-derived behavior. Runtime, activation, telemetry, persistence, update and uninstall behavior are not observed.")
    doc.add_paragraph(f"Verification procedure: provision a disposable {platforms} host; install from the approved hash-locked repository; enable IDE and endpoint logging; test file-root escapes, symlinks, network shares where applicable, data URIs, private IPs, DNS rebinding, redirects, timeouts, large inputs, nested archives, cancellation, malformed HTTP and uninstall cleanup. Record evidence and restore the host.")


def build():
    make_figure()
    stage_files = sorted(STAGES.glob("[0-9][0-9]-*.md"))
    manifest = json.loads((STAGES / "manifest.json").read_text())
    doc = Document()
    configure(doc)
    add_title_page(doc)
    doc.add_heading("Document Control", level=1)
    add_table(doc, ["Field", "Value"], [
        ("Document", "Microsoft MarkItDown MCP Security Assessment"),
        ("Assessed product", "Microsoft-published markitdown-mcp 0.0.1a4 with markitdown 0.1.6"),
        ("Scope", "VS Code: Windows 11 and macOS; Visual Studio: Windows 11"),
        ("Business context", "UK financial services; internal-confidential information"),
        ("Assessment date", "25 July 2026"),
        ("Decision", "Do not approve for production"),
        ("Runtime status", "Selected and permitted, but not executed: no disposable representative host"),
        ("Authoritative status", "Authoritative after automated structural and Word visual QA"),
    ])
    doc.add_heading("Revision History", level=1)
    add_table(doc, ["Version", "Date", "Status", "Change"], [("1.0", "25 July 2026", "Authoritative", "Initial evidence-led assessment")])
    doc.add_heading("Contents", level=1)
    doc.add_paragraph("Use Word’s Navigation pane for the structured contents. The heading hierarchy is authoritative; internal REF links navigate to the reference register.")
    add_bullets(doc, ["Executive Summary", "Purpose and Function Overview", "Scope, Assumptions and Methodology", "Product Identity and Version", "Part I – VS Code", "Part II – Visual Studio", "Part III – Installed Agent Skills", "Consolidated Risk Register", "References"])
    doc.add_heading("Executive Summary", level=1)
    add_bullets(doc, [
        "Decision — Do not approve the published MarkItDown MCP 0.0.1a4 for production or regulated internal-confidential workloads in the assessed IDEs.",
        "Critical authority — The single MCP tool accepts arbitrary HTTP, HTTPS, file and data URIs, creating direct local-file disclosure, server-side request forgery and content-ingestion paths without an intrinsic root or host allowlist.",
        "Transport exposure — Published HTTP and SSE operation is unauthenticated and resolves to MCP Python SDK 1.8.1, which is affected by applicable DNS-rebinding and denial-of-service advisories.",
        "Evidence quality — Package identity, hashes, source correspondence, static analysis, malware scans and dependency resolution were verified; runtime behavior was not observed because no disposable representative Windows or macOS host was available.",
        "Conditional future path — Reassessment may consider a remediated, pinned STDIO-only build inside a least-privilege IDE sandbox after URI policy, resource bounds, update governance, telemetry checks and representative runtime tests pass.",
    ])
    p = doc.add_paragraph("The recommendation is evidence-based and does not rely on scanner absence as proof of safety. The official project and publication records establish Microsoft provenance for the core project and packages")
    cite(p, 1); cite(p, 2); cite(p, 3)
    doc.add_paragraph("The highest risks arise from intended product authority and transport defaults, not from detected malicious code. For a UK financial-services environment, that distinction is decisive: a benign utility can still breach confidentiality, operational resilience and change-control expectations when invoked by an agent against attacker-influenced inputs.")
    doc.add_heading("Purpose Function Overview", level=1)
    doc.add_paragraph("MarkItDown converts supported documents, archives, images and network resources to Markdown so an AI client can consume their textual content. The assessed product is the Microsoft-published Python MCP server, not an unrelated marketplace VSIX. The server exposes convert_to_markdown(uri), constructs MarkItDown with its published defaults, and returns converted text to the MCP client. This creates a compact interface with unusually broad authority because the URI determines both source location and processing path.")
    add_table(doc, ["Capability", "Intended value", "Security consequence"], [
        ("Local conversion", "Normalize documents for agent use", "Reads content reachable by the server process"),
        ("Remote conversion", "Fetch and normalize URLs", "Creates outbound request and SSRF exposure"),
        ("Archive handling", "Extract and convert container content", "Expands parser, recursion and resource-exhaustion surface"),
        ("MCP tool exposure", "One callable function for IDE agents", "Delegates URI choice to model-mediated workflows"),
    ])
    doc.add_heading("Scope Assumptions Methodology", level=1)
    add_narrative(doc, "Assessment Scope", [
        "In scope are VS Code on Windows 11 and macOS, Visual Studio on Windows 11, the published markitdown-mcp 0.0.1a4 distribution, the resolved markitdown 0.1.6 core and the dependency sets produced for Python 3.12. The data classification is internal-confidential in a UK financial-services organization.",
        "Dynamic analysis, malware analysis and static analysis were authorized. Dynamic analysis was selected but fail-closed because the assessment did not have a disposable or revertible representative Windows host and did not establish a disposable macOS host. No third-party package was executed on the analyst workstation.",
    ])
    add_narrative(doc, "Evidence Method", [
        "The method separated publisher identity, package acquisition, cryptographic verification, archive inspection, source comparison, static security analysis, malware scanning, dependency resolution, advisory correlation, architecture analysis, privacy analysis and IDE-specific control assessment. Claims are labelled verified, inferred, not observed, not applicable or unknown.",
        "Static inspection used source review, Semgrep and YARA. Malware inspection used current ClamAV signatures and manual triage of heuristic YARA matches. Supply-chain work generated platform-specific hashed dependency resolutions and correlated direct dependencies with OSV advisories. Runtime claims remain limitations rather than silently becoming assumptions.",
    ])
    add_table(doc, ["Analysis", "Selected", "Executed", "Evidence status"], [
        ("Static analysis", "Yes", "Yes", "Verified"), ("Malware analysis", "Yes", "Yes", "Verified; no malware detected"),
        ("Dynamic analysis", "Yes", "No", "Not observed; disposable environment unavailable"), ("Network behavior", "Yes", "Static only", "Inferred; runtime verification required"),
    ])
    doc.add_page_break()
    doc.add_heading("Product Identity and Version", level=1)
    p = doc.add_paragraph("Official identity was resolved to the Microsoft MarkItDown repository and Microsoft-published PyPI artifacts")
    cite(p, 1); cite(p, 2); cite(p, 3)
    add_table(doc, ["Component", "Version", "Digest / source", "Identity conclusion"], [
        ("markitdown-mcp wheel", "0.0.1a4 pre-release", "SHA-256 7fb06fff7d722ec108d08752704dc8f313f7fd267e00dff546131ec229645230", "Microsoft project; MCP-specific source tag absent"),
        ("markitdown-mcp sdist", "0.0.1a4 pre-release", "SHA-256 309c94dc883311e6909d849382a6c7bc402dfb2692dab448c136c6864c6bf49e", "Published hash matched"),
        ("markitdown wheel", "0.1.6", "SHA-256 07b2d5bf87b5c53e13a9f2fdc440df8ccc85e26f40c1e557781727b700049775", "Matched tagged source"),
        ("markitdown source", "v0.1.6", "commit e144e0a2be95b34df17433bac904e635f2c5e551", "Tag correspondence verified"),
    ])
    doc.add_paragraph("The package under assessment is not a Visual Studio or VS Code extension package. It is a local MCP server configured through IDE MCP support. The similarly named third-party marketplace extension is excluded.")
    doc.add_heading("Architecture and Trust-Boundary Overview", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_run = p.add_run()
    picture_run.add_picture(str(FIG), width=Inches(7.2))
    for drawing_property in picture_run._r.xpath(".//wp:docPr"):
        drawing_property.set("descr", "Trust-boundary diagram showing untrusted input flowing through the IDE and MCP transport into MarkItDown, then to local files or network resources.")
        drawing_property.set("title", "MarkItDown MCP trust boundaries")
    doc.add_paragraph("Figure 1. The IDE and agent cross into a local MCP process; URI handling then crosses into filesystem, archive parser and network trust zones. Returned Markdown crosses back into model context. Controls must constrain both directions.")
    ide_part(doc, "Part I – VS Code Assessment", "VS Code", "Windows 11 and macOS", 4)
    ide_part(doc, "Part II – Visual Studio Assessment", "Visual Studio", "Windows 11", 5)
    doc.add_heading("Part III – Installed Agent Skills", level=1)
    doc.add_paragraph("Not applicable as a shipped component: the published wheels contain no installed Agent Skill definition. The product is an MCP server and Python conversion library. Installed enterprise or repository Agent Skills remain an adjacent influence surface because their instructions can cause tool invocation.")
    add_table(doc, ["Question", "Disposition", "Control"], [
        ("Bundled Agent Skill?", "Not applicable / not present", "Verify future packages and SBOMs"),
        ("Can external instructions invoke MCP?", "Yes, inferred from agent workflow", "Treat instructions as untrusted; require approval"),
        ("Skill persistence assessed?", "Not applicable to package", "Assess separately where skills are installed"),
        ("Prompt injection relevance", "Applicable", "Content isolation, provenance, confirmation and output handling"),
    ])
    doc.add_heading("Cross-IDE Comparison", level=1)
    add_table(doc, ["Dimension", "VS Code Windows 11", "VS Code macOS", "Visual Studio Windows 11"], [
        ("Packaging", "Python MCP subprocess", "Python MCP subprocess", "Python MCP subprocess"), ("Preferred transport", "STDIO only", "STDIO only", "STDIO only"),
        ("HTTP/SSE", "Prohibited", "Prohibited", "Prohibited"), ("Primary boundary", "Windows endpoint policy", "macOS endpoint policy", "Windows endpoint policy"),
        ("Runtime evidence", "Not observed", "Not observed", "Not observed"), ("Decision", "Do not approve production", "Do not approve production", "Do not approve production"),
    ])
    doc.add_heading("Consolidated Supply Chain Assessment", level=1)
    p = doc.add_paragraph("Official repository and PyPI records support identity, while exact MCP source correspondence remains incomplete")
    cite(p, 1); cite(p, 2); cite(p, 3)
    add_table(doc, ["Evidence", "Result", "Interpretation"], [
        ("Archive safety", "No absolute paths, traversal entries, symlinks or native binaries", "Verified hygiene; not approval"),
        ("ClamAV 1.5.3 / database 28071", "46 files; zero infected", "No known malware signature detected"),
        ("YARA 4.5.5 baseline", "Three heuristic matches", "Manually triaged as benign-explainable"),
        ("Semgrep 1.171.0", "1,066 Python rules; 41 files; zero findings", "No rule finding; design findings remain"),
        ("Trivy 0.72.0", "Low Dockerfile healthcheck issue", "Docker path not assessed IDE install"),
        ("OSV Scanner 2.4.0", "Three applicable advisories in mcp 1.8.1", "Upgrade constraint requires product change"),
    ])
    p = doc.add_paragraph("Applicable advisories cover DNS rebinding and two denial-of-service conditions")
    cite(p, 6); cite(p, 7); cite(p, 8)
    doc.add_heading("Consolidated Privacy and Data Protection Assessment", level=1)
    doc.add_paragraph("The conversion function processes content rather than merely metadata. Local files, remote resources and archives may contain personal data, special-category data, customer financial information, authentication material and internal-confidential business records. Returned Markdown can propagate that content into IDE chat, model processing, logs and downstream copy and paste.")
    p = doc.add_paragraph("UK governance should align with ICO AI guidance, the Data Protection Act, operational resilience expectations and third-party risk management")
    cite(p, 17); cite(p, 18); cite(p, 19); cite(p, 20)
    add_table(doc, ["Principle", "Risk", "Required control"], [
        ("Lawfulness and transparency", "Unclear model processing path", "Approved use case and DPIA trigger assessment"), ("Purpose limitation", "General URI exceeds purpose", "Enforced source policy"),
        ("Data minimisation", "Whole documents returned", "Selective extraction, redaction and output caps"), ("Accuracy", "Conversion omits context", "Human verification"),
        ("Storage limitation", "IDE/model retention unknown", "Configured retention and deletion test"), ("Integrity and confidentiality", "File/network disclosure", "Least privilege, DLP and egress control"),
    ])
    doc.add_page_break()
    doc.add_heading("Enterprise Controls Roadmap", level=1)
    add_table(doc, ["Milestone / target date", "Owner", "Action", "Verification test", "Control strength"], [
        ("0–30 days", "Product Security", "Block HTTP/SSE and package 0.0.1a4 in production channels", "Policy tests deny install/configuration", "Strong preventive"),
        ("0–30 days", "Engineering", "Define STDIO wrapper with file and network policy", "Adversarial tests fail closed", "Strong preventive"),
        ("0–30 days", "Privacy / Legal", "Complete data-flow and DPIA decision", "Signed decision", "Governance"),
        ("31–90 days", "Supply Chain", "Rebuild with fixed SDK, hashes, SBOM and provenance", "Pipeline verifies evidence", "Strong preventive"),
        ("31–90 days", "Endpoint Security", "Create process, socket, egress and DLP detections", "Purple-team alerts", "Detective"),
        ("31–90 days", "IDE Platform", "Run disposable-host verification", "Acceptance evidence passes", "Verification"),
        ("Before pilot", "Risk Owner", "Review residual risks and narrow pilot", "Signed decision and expiry", "Governance"),
    ])
    doc.add_heading("Detection Opportunities and Monitoring Plan", level=1)
    add_table(doc, ["Signal", "Detection opportunity", "Response"], [
        ("Process", "IDE starts unapproved Python or markitdown command", "Alert and terminate"), ("Listener", "MCP process binds TCP or uses HTTP/SSE", "Critical alert"),
        ("Network", "Private, link-local or metadata request", "Block and investigate"), ("Filesystem", "Read outside approved roots", "Block and investigate"),
        ("Resource", "High memory/CPU or archive expansion", "Cancel and isolate"), ("Content", "DLP match in returned Markdown", "Prevent transfer"),
        ("Supply chain", "Hash, publisher, SBOM or advisory drift", "Quarantine and reassess"),
    ])
    doc.add_heading("Consolidated Risk Register", level=1)
    doc.add_paragraph("Scoring legend: 1–4 Low; 5–9 Moderate; 10–16 High; 17–25 Critical. Score = likelihood × impact. Residual score assumes only stated controls. Treatment must be accepted by the named owner.")
    risks = [
        ("F-001", "Unrestricted URI authority: local-file disclosure and SSRF", "4×5=20 Critical", "4×4=16 High", "Engineering", "Mitigate"),
        ("F-002", "Unauthenticated vulnerable HTTP/SSE transport", "4×5=20 Critical", "4×5=20 Critical", "Product Security", "Avoid"),
        ("F-003", "Unbounded network, data and archive processing", "4×4=16 High", "4×4=16 High", "Engineering", "Mitigate"),
        ("F-004", "Broad, pre-release and vulnerable dependency chain", "4×4=16 High", "3×4=12 High", "Supply Chain", "Mitigate"),
        ("F-005", "Runtime and lifecycle unverified", "3×4=12 High", "3×4=12 High", "IDE Platform", "Defer"),
        ("F-006", "Confidential content enters model context", "4×5=20 Critical", "3×5=15 High", "Data Owner", "Mitigate"),
        ("F-007", "MCP source-to-package provenance gap", "3×3=9 Moderate", "3×3=9 Moderate", "Supply Chain", "Mitigate"),
    ]
    add_table(doc, ["ID", "Risk", "Inherent", "Residual", "Owner", "Treatment"], risks)
    doc.add_heading("Residual Risk and Approval Recommendation", level=1)
    doc.add_paragraph("Recommendation: Do not approve for production, regulated processing or internal-confidential content. Do not enable published HTTP, SSE or Streamable HTTP transports. The residual profile includes one Critical and multiple High risks, and runtime evidence is absent across every requested host.")
    doc.add_paragraph("A risk owner may sponsor reassessment of a materially remediated build. Entry criteria are: fixed MCP SDK; STDIO-only transport; enforced URI and filesystem policy; denied private-address access; redirect revalidation; input, output, archive, recursion, time and concurrency bounds; immutable dependency locks; internal promotion; complete SBOM and provenance; privacy approval; detections; and successful disposable-host verification.")
    doc.add_heading("Limitations and Confidence Levels", level=1)
    add_table(doc, ["Area", "State", "Confidence / limitation"], [
        ("Publisher and package identity", "Verified", "High"), ("Hashes and archive contents", "Verified", "High"), ("Core source correspondence", "Verified", "High"),
        ("MCP source correspondence", "Unknown", "Medium; no scoped tag or attestation"), ("Static analysis and malware scan", "Verified", "High within tool limits"),
        ("Runtime and lifecycle", "Not observed", "No representative disposable host"), ("Agent Skills", "Not applicable as bundled component", "High from inspection"),
    ])
    doc.add_heading("Evidence Register", level=1)
    evidence_rows = []
    for idx, path in enumerate(stage_files, 1):
        text = path.read_text()
        state = re.search(r"Evidence state:\s*([^\n]+)", text)
        evidence_rows.append((f"EV-{idx:03d}", path.name, state.group(1).strip() if state else "Recorded", hashlib.sha256(text.encode()).hexdigest()[:20], "Retained locally"))
    add_table(doc, ["Evidence ID", "Stage", "State", "SHA-256 prefix", "Retention"], evidence_rows)
    doc.add_paragraph(f"Stage manifest target: {manifest.get('target', 'Microsoft MarkItDown')}. All fifteen stage records were ingested into the report build. Raw artifacts and scanner caches are not embedded.")
    doc.add_heading("References", level=1)
    for i, (label, name, url) in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        add_bookmark(p, label.replace("-", "_"), 100 + i)
        p.add_run(f"{label} — {name}. ")
        add_hyperlink(p, f"Open {name}", url=url)
        p.add_run(" Accessed 25 July 2026.")
    doc.add_heading("Appendices", level=1)
    doc.add_heading("Appendix A – Analysis Tooling and Results", level=2)
    add_table(doc, ["Tool", "Version / basis", "Result"], [
        ("ClamAV", "1.5.3; database 28071", "46 files; zero infected"), ("Semgrep", "1.171.0; 1,066 Python rules", "41 files; zero findings"),
        ("YARA", "4.5.5; baseline rules", "Three benign-explainable matches"), ("OSV Scanner", "2.4.0", "Three applicable mcp advisories"),
        ("Syft / Grype", "1.49.0 / bounded SBOM", "Components enumerated; zero Grype matches"), ("Trivy", "0.72.0", "Low Dockerfile healthcheck issue"),
    ])
    doc.add_heading("Appendix B – Runtime Verification Protocol", level=2)
    add_bullets(doc, [
        "Provision disposable Windows 11 hosts for VS Code and Visual Studio and a disposable macOS host for VS Code; snapshot before installation.",
        "Install exact approved artifacts into dedicated Python environments from a controlled index; verify hashes and resolved SBOM.",
        "Capture process, child-process, filesystem, DNS, proxy, socket and HTTP behavior without real confidential information.",
        "Exercise traversal, symlink, network-share, private-address, redirect, data URI, oversized response, decompression bomb, nested archive, malformed request, cancellation and concurrency cases.",
        "Test update, rollback, configuration precedence, workspace trust, consent, persistence and uninstall residue. Restore snapshots and preserve sanitized evidence.",
    ])
    doc.add_heading("Appendix C – Framework Reference Notes", level=2)
    for n in range(9, 17):
        p = doc.add_paragraph("Framework basis and control design reference")
        cite(p, n)
    doc.add_paragraph("Framework mappings are control-design aids, not claims of certification or complete compliance. OWASP, NIST, MITRE, NCSC and CIS sources structure threats and controls; UK regulatory sources inform governance and resilience.")
    doc.add_heading("Glossary", level=1)
    add_table(doc, ["Term", "Meaning"], [
        ("Agent Skill", "Instruction package distinct from an MCP server"), ("MCP", "Model Context Protocol"), ("STDIO", "Process-local standard input/output transport"),
        ("SSE", "Server-Sent Events network transport"), ("SSRF", "Server-side request forgery"), ("SBOM", "Software bill of materials"), ("DPIA", "Data protection impact assessment"),
        ("Verified", "Supported by inspected or executed evidence"), ("Inferred", "Reasoned without runtime observation"), ("Not observed", "Not witnessed in representative runtime"),
        ("Unknown", "Evidence insufficient to conclude"),
    ])
    doc.add_paragraph("END OF AUTHORITATIVE REPORT")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
