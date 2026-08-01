#!/usr/bin/env python3
"""Build an authoritative assessment DOCX from a hash-bound report model."""

from __future__ import annotations

import argparse
import errno
import hashlib
import importlib.metadata
import ipaddress
import json
import os
import re
import socket
import stat
import sys
import tempfile
from contextlib import contextmanager
from datetime import UTC, datetime
from pathlib import Path
from typing import BinaryIO, Iterator
from urllib.parse import urlsplit

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from portable_fs import is_link_or_reparse, require_real_directory
from PIL import Image, ImageDraw, ImageFont
from safe_stage_inputs import consume_validated_stages, read_bounded_regular_file

RUN_KEY = re.compile(r"^\d{4}-\d{2}-\d{2}-v\d+\.\d+$")
IDENTIFIER = re.compile(r"^[a-z][a-z0-9_.-]{2,127}$")
REFERENCE_ID = re.compile(r"^REF-\d{3}$")
EVIDENCE_ID = re.compile(r"^EVD-[A-Z0-9-]{3,64}$")
FINDING_ID = re.compile(r"^(?:F|RISK)-\d{3}$")
CITATION = re.compile(r"\bREF-(\d{3})\b")
MAX_MODEL_BYTES = 4 * 1024 * 1024
MAX_STRING_BYTES = 64 * 1024
MAX_ITEMS = 500
PUBLIC_HOST_SUFFIXES_DENY = (".internal", ".local", ".localhost")
PROHIBITED_REPORT_TEXT = (
    re.compile(r"\bLL-\d{4}\b", re.IGNORECASE),
    re.compile(r"\bcentral lessons(?:-learned)?\b", re.IGNORECASE),
    re.compile(r"\bRCA[- ]\d+\b", re.IGNORECASE),
    re.compile(r"/Users/[^/\s]+/", re.IGNORECASE),
    re.compile(r"file://", re.IGNORECASE),
)
DECISIONS = {
    "Approve",
    "Approve with conditions",
    "Defer pending evidence",
    "Do not approve",
}
RATINGS = {"Low", "Moderate", "High", "Critical"}
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


class ModelError(ValueError):
    """Raised when a report model is unsafe or incomplete."""


FileIdentity = tuple[int, int]


def _identity(metadata: os.stat_result) -> FileIdentity:
    return metadata.st_dev, metadata.st_ino


def _require_regular_identity(path: Path, expected: FileIdentity, label: str) -> None:
    try:
        metadata = path.lstat()
    except FileNotFoundError as exc:
        raise ModelError(f"{label} disappeared during generation") from exc
    if is_link_or_reparse(metadata) or not stat.S_ISREG(metadata.st_mode):
        raise ModelError(f"{label} is no longer a regular non-symlink file")
    if _identity(metadata) != expected:
        raise ModelError(f"{label} was replaced during generation")


def _resolve_workspace_root(candidate: Path) -> Path:
    """Resolve a caller-selected workspace without accepting a symlink root."""

    absolute = Path(os.path.abspath(os.fspath(candidate)))
    try:
        metadata = absolute.lstat()
    except FileNotFoundError as exc:
        raise ModelError("workspace root does not exist") from exc
    if is_link_or_reparse(metadata) or not stat.S_ISDIR(metadata.st_mode):
        raise ModelError("workspace root must be a real non-symlink directory")
    return absolute.resolve(strict=True)


def _require_within_workspace(path: Path, workspace_root: Path, label: str) -> None:
    try:
        path.relative_to(workspace_root)
    except ValueError as exc:
        raise ModelError(f"{label} must remain inside the workspace root") from exc


def _resolve_stage_root(candidate: Path, workspace_root: Path) -> Path:
    absolute = Path(os.path.abspath(os.fspath(candidate)))
    try:
        metadata = absolute.lstat()
    except FileNotFoundError as exc:
        raise ModelError("stage root does not exist") from exc
    if is_link_or_reparse(metadata) or not stat.S_ISDIR(metadata.st_mode):
        raise ModelError("stage root must be a real non-symlink directory")
    resolved = absolute.resolve(strict=True)
    _require_within_workspace(resolved, workspace_root, "stage root")
    return resolved


def _resolve_destination(candidate: Path, label: str, workspace_root: Path) -> Path:
    """Resolve only the destination parent so a leaf symlink remains observable."""

    absolute = Path(os.path.abspath(os.fspath(candidate)))
    if absolute.name in {"", ".", ".."}:
        raise ModelError(f"{label} must name a file")
    try:
        parent = require_real_directory(absolute.parent)
    except (FileNotFoundError, ValueError) as exc:
        raise ModelError(f"{label} parent directory does not exist") from exc
    if not parent.is_dir():
        raise ModelError(f"{label} parent must be a directory")
    try:
        parent.relative_to(workspace_root)
    except ValueError as exc:
        raise ModelError(f"{label} must remain inside the workspace root") from exc
    destination = parent / absolute.name
    try:
        metadata = destination.lstat()
    except FileNotFoundError:
        return destination
    if is_link_or_reparse(metadata):
        raise ModelError(f"{label} must not be a symlink, junction, or reparse point")
    raise ModelError(f"{label} already exists; choose a new versioned path")


def _create_owned_file(path: Path, label: str) -> FileIdentity:
    flags = os.O_CREAT | os.O_EXCL | os.O_RDWR | getattr(os, "O_BINARY", 0) | getattr(os, "O_NOFOLLOW", 0)
    try:
        descriptor = os.open(path, flags, 0o600)
    except FileExistsError as exc:
        raise ModelError(f"{label} already exists; refusing to overwrite it") from exc
    try:
        metadata = os.fstat(descriptor)
        if not stat.S_ISREG(metadata.st_mode):
            raise ModelError(f"{label} is not a regular file")
        return _identity(metadata)
    finally:
        os.close(descriptor)


def _create_secure_temp(destination: Path, label: str) -> tuple[Path, FileIdentity]:
    descriptor, filename = tempfile.mkstemp(
        prefix=f".{destination.name}.",
        suffix=f".tmp{destination.suffix}",
        dir=destination.parent,
    )
    path = Path(filename)
    try:
        metadata = os.fstat(descriptor)
        if not stat.S_ISREG(metadata.st_mode):
            raise ModelError(f"temporary {label} is not a regular file")
        return path, _identity(metadata)
    finally:
        os.close(descriptor)


@contextmanager
def _open_owned_file(
    path: Path,
    expected: FileIdentity,
    label: str,
    *,
    write: bool,
) -> Iterator[BinaryIO]:
    flags = (os.O_WRONLY | os.O_TRUNC) if write else os.O_RDONLY
    flags |= getattr(os, "O_BINARY", 0) | getattr(os, "O_NOFOLLOW", 0)
    descriptor = os.open(path, flags)
    try:
        metadata = os.fstat(descriptor)
        if not stat.S_ISREG(metadata.st_mode) or _identity(metadata) != expected:
            raise ModelError(f"{label} was replaced during generation")
        mode = "wb" if write else "rb"
        with os.fdopen(descriptor, mode) as stream:
            descriptor = -1
            yield stream
            if write:
                stream.flush()
                os.fsync(stream.fileno())
    finally:
        if descriptor >= 0:
            os.close(descriptor)


def _unlink_if_owned(path: Path, expected: FileIdentity) -> bool:
    """Best-effort cleanup that preserves a path replaced by another actor."""

    try:
        metadata = path.lstat()
    except FileNotFoundError:
        return False
    if (
        is_link_or_reparse(metadata)
        or not stat.S_ISREG(metadata.st_mode)
        or _identity(metadata) != expected
    ):
        return False
    path.unlink()
    return True


def _sha256_owned_file(path: Path, expected: FileIdentity, label: str) -> str:
    digest = hashlib.sha256()
    with _open_owned_file(path, expected, label, write=False) as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    _require_regular_identity(path, expected, label)
    return digest.hexdigest()


def _publish_exclusive(
    temporary: Path,
    destination: Path,
    expected: FileIdentity,
    label: str,
) -> None:
    _require_regular_identity(temporary, expected, f"temporary {label}")
    try:
        os.link(temporary, destination, follow_symlinks=False)
    except FileExistsError as exc:
        raise ModelError(f"{label} appeared during generation; refusing to overwrite it") from exc
    except OSError as exc:
        if exc.errno == errno.EEXIST:
            raise ModelError(
                f"{label} appeared during generation; refusing to overwrite it"
            ) from exc
        raise
    _require_regular_identity(destination, expected, label)


def _publish_bundle(
    artifacts: list[tuple[Path, Path, FileIdentity, str]],
) -> None:
    published: list[tuple[Path, FileIdentity]] = []
    try:
        for temporary, destination, expected, label in artifacts:
            _publish_exclusive(temporary, destination, expected, label)
            published.append((destination, expected))
    except Exception:
        for destination, expected in reversed(published):
            _unlink_if_owned(destination, expected)
        raise


def _text(value: object, field: str, *, allow_empty: bool = False) -> str:
    if not isinstance(value, str) or (not allow_empty and not value.strip()):
        raise ModelError(f"{field} must be a non-empty string")
    if "\x00" in value or len(value.encode("utf-8")) > MAX_STRING_BYTES:
        raise ModelError(f"{field} is oversized or contains NUL")
    for pattern in PROHIBITED_REPORT_TEXT:
        if pattern.search(value):
            raise ModelError(f"{field} contains prohibited internal or unsafe text")
    return value.strip()


def _list(value: object, field: str, *, minimum: int = 0, maximum: int = MAX_ITEMS) -> list[object]:
    if not isinstance(value, list) or not minimum <= len(value) <= maximum:
        raise ModelError(f"{field} must contain between {minimum} and {maximum} items")
    return value


def _public_https(value: object, field: str) -> str:
    url = _text(value, field)
    parsed = urlsplit(url)
    if (
        parsed.scheme.casefold() != "https"
        or not parsed.hostname
        or parsed.username is not None
        or parsed.password is not None
    ):
        raise ModelError(f"{field} must be a public HTTPS URL without credentials")
    hostname = parsed.hostname.casefold().rstrip(".")
    if hostname == "localhost" or hostname.endswith(PUBLIC_HOST_SUFFIXES_DENY):
        raise ModelError(f"{field} targets a private hostname")
    try:
        address = ipaddress.ip_address(hostname)
    except ValueError:
        # inet_aton deliberately recognises legacy IPv4 spellings such as
        # 127.1, 2130706433 and 0x7f000001. Treat them as addresses instead of
        # allowing URL parsers or downstream clients to reinterpret them.
        try:
            address = ipaddress.ip_address(socket.inet_aton(hostname))
        except OSError:
            return url
    if (
        address.is_private
        or address.is_loopback
        or address.is_link_local
        or address.is_multicast
        or address.is_reserved
        or address.is_unspecified
    ):
        raise ModelError(f"{field} targets a non-public address")
    return url


def _validate_table(table: object, field: str) -> dict[str, object]:
    if not isinstance(table, dict) or set(table) != {"title", "columns", "rows"}:
        raise ModelError(f"{field} has missing or unexpected fields")
    title = _text(table["title"], f"{field}.title")
    columns = [_text(item, f"{field}.columns") for item in _list(table["columns"], f"{field}.columns", minimum=1, maximum=12)]
    rows: list[list[str]] = []
    for row_index, raw_row in enumerate(_list(table["rows"], f"{field}.rows", minimum=1, maximum=250)):
        values = [_text(item, f"{field}.rows[{row_index}]") for item in _list(raw_row, f"{field}.rows[{row_index}]", minimum=len(columns), maximum=len(columns))]
        rows.append(values)
    return {"title": title, "columns": columns, "rows": rows}


def validate_report_model(raw: object) -> dict[str, object]:
    if not isinstance(raw, dict):
        raise ModelError("report model must be a JSON object")
    required = {
        "schema_version", "assessment", "target", "publisher", "extension_id",
        "version", "run_key", "assessment_date", "document_version",
        "classification", "decision", "overall_residual_risk", "review_trigger",
        "ide_scope", "executive_outcomes", "approval_conditions", "sections",
        "findings", "evidence", "references", "glossary", "figure",
        "derivative_sources",
    }
    optional = {"revision_history"}
    if not required <= set(raw) or set(raw) - required - optional:
        missing = sorted(required - set(raw))
        unexpected = sorted(set(raw) - required - optional)
        raise ModelError(f"report model fields differ; missing={missing}, unexpected={unexpected}")
    if raw["schema_version"] != 2:
        raise ModelError("unsupported report model schema")

    model: dict[str, object] = {"schema_version": 2}
    for field in (
        "assessment", "target", "publisher", "extension_id", "version",
        "assessment_date", "document_version", "review_trigger",
    ):
        model[field] = _text(raw[field], field)
    run_key = _text(raw["run_key"], "run_key")
    if not RUN_KEY.fullmatch(run_key):
        raise ModelError("run_key is invalid")
    model["run_key"] = run_key
    if raw["classification"] != "PUBLIC":
        raise ModelError("new authoritative reports must be PUBLIC classified")
    model["classification"] = "PUBLIC"
    if raw["decision"] not in DECISIONS:
        raise ModelError("decision is invalid")
    model["decision"] = raw["decision"]
    if raw["overall_residual_risk"] not in RATINGS:
        raise ModelError("overall_residual_risk is invalid")
    model["overall_residual_risk"] = raw["overall_residual_risk"]
    model["ide_scope"] = [
        _text(item, "ide_scope")
        for item in _list(raw["ide_scope"], "ide_scope", minimum=1, maximum=4)
    ]
    model["executive_outcomes"] = [
        _text(item, "executive_outcomes")
        for item in _list(raw["executive_outcomes"], "executive_outcomes", minimum=5, maximum=10)
    ]
    model["approval_conditions"] = [
        _text(item, "approval_conditions")
        for item in _list(raw["approval_conditions"], "approval_conditions", minimum=1, maximum=20)
    ]
    default_revision = [{
        "version": model["document_version"],
        "date": model["assessment_date"],
        "status": "Candidate pending native Word closeout",
        "change": "Initial schema-driven assessment",
    }]
    raw_revisions = raw.get("revision_history", default_revision)
    revisions: list[dict[str, str]] = []
    revision_fields = ("version", "date", "status", "change")
    seen_revision_versions: set[str] = set()
    previous_revision_date = None
    previous_revision_version: tuple[int, int] | None = None
    for index, revision in enumerate(
        _list(raw_revisions, "revision_history", minimum=1, maximum=20)
    ):
        if not isinstance(revision, dict) or set(revision) != set(revision_fields):
            raise ModelError(
                f"revision_history[{index}] has missing or unexpected fields"
            )
        converted: dict[str, str] = {}
        for field in revision_fields:
            value = _text(revision[field], f"revision_history[{index}].{field}")
            if len(value.encode("utf-8")) > 4096:
                raise ModelError(
                    f"revision_history[{index}].{field} exceeds 4096 bytes"
                )
            converted[field] = value
        version = converted["version"]
        if not re.fullmatch(r"\d+\.\d+", version) or version in seen_revision_versions:
            raise ModelError(
                f"revision_history[{index}].version is invalid or duplicate"
            )
        seen_revision_versions.add(version)
        revision_version = tuple(int(part) for part in version.split("."))
        if (
            previous_revision_version is not None
            and revision_version <= previous_revision_version
        ):
            raise ModelError("revision_history versions must be in increasing order")
        previous_revision_version = revision_version
        if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", converted["date"]):
            raise ModelError(
                f"revision_history[{index}].date must be a valid ISO date"
            )
        try:
            revision_date = datetime.strptime(converted["date"], "%Y-%m-%d").date()
        except ValueError as exc:
            raise ModelError(
                f"revision_history[{index}].date must be a valid ISO date"
            ) from exc
        if previous_revision_date is not None and revision_date < previous_revision_date:
            raise ModelError("revision_history dates must be in nondecreasing order")
        previous_revision_date = revision_date
        revisions.append(converted)
    if revisions[-1]["version"] != model["document_version"]:
        raise ModelError("latest revision history version must match document_version")
    model["revision_history"] = revisions

    sections: list[dict[str, object]] = []
    seen_section_ids: set[str] = set()
    for index, section in enumerate(_list(raw["sections"], "sections", minimum=20, maximum=100)):
        if not isinstance(section, dict) or set(section) != {"id", "heading", "level", "paragraphs", "bullets", "tables"}:
            raise ModelError(f"sections[{index}] has missing or unexpected fields")
        section_id = _text(section["id"], f"sections[{index}].id")
        if not IDENTIFIER.fullmatch(section_id) or section_id in seen_section_ids:
            raise ModelError(f"invalid or duplicate section id: {section_id}")
        seen_section_ids.add(section_id)
        level = section["level"]
        if not isinstance(level, int) or isinstance(level, bool) or level not in (1, 2, 3):
            raise ModelError(f"sections[{index}].level is invalid")
        sections.append({
            "id": section_id,
            "heading": _text(section["heading"], f"sections[{index}].heading"),
            "level": level,
            "paragraphs": [_text(item, f"sections[{index}].paragraphs") for item in _list(section["paragraphs"], f"sections[{index}].paragraphs", maximum=30)],
            "bullets": [_text(item, f"sections[{index}].bullets") for item in _list(section["bullets"], f"sections[{index}].bullets", maximum=30)],
            "tables": [_validate_table(item, f"sections[{index}].tables") for item in _list(section["tables"], f"sections[{index}].tables", maximum=12)],
        })
    model["sections"] = sections

    findings: list[dict[str, str]] = []
    finding_fields = {
        "id", "title", "scope", "scenario", "evidence_ids", "likelihood", "impact",
        "inherent", "controls", "control_strength", "residual_likelihood",
        "residual_impact", "residual", "recommendation", "owner", "priority",
        "target_date", "verification", "mappings", "confidence",
    }
    seen_findings: set[str] = set()
    for index, finding in enumerate(_list(raw["findings"], "findings", minimum=1, maximum=100)):
        if not isinstance(finding, dict) or set(finding) != finding_fields:
            raise ModelError(f"findings[{index}] has missing or unexpected fields")
        converted = {field: _text(finding[field], f"findings[{index}].{field}") for field in finding_fields}
        if not FINDING_ID.fullmatch(converted["id"]) or converted["id"] in seen_findings:
            raise ModelError(f"invalid or duplicate finding id: {converted['id']}")
        seen_findings.add(converted["id"])
        findings.append(converted)
    model["findings"] = findings

    evidence: list[dict[str, str]] = []
    seen_evidence: set[str] = set()
    for index, item in enumerate(_list(raw["evidence"], "evidence", minimum=5, maximum=250)):
        fields = {"id", "title", "source", "method", "state", "limitation"}
        if not isinstance(item, dict) or set(item) != fields:
            raise ModelError(f"evidence[{index}] has missing or unexpected fields")
        converted = {field: _text(item[field], f"evidence[{index}].{field}") for field in fields}
        if not EVIDENCE_ID.fullmatch(converted["id"]) or converted["id"] in seen_evidence:
            raise ModelError(f"invalid or duplicate evidence id: {converted['id']}")
        seen_evidence.add(converted["id"])
        evidence.append(converted)
    model["evidence"] = evidence

    references: list[dict[str, str]] = []
    seen_references: set[str] = set()
    for index, item in enumerate(_list(raw["references"], "references", minimum=12, maximum=250)):
        fields = {"id", "title", "publisher", "url", "accessed", "applicability"}
        if not isinstance(item, dict) or set(item) != fields:
            raise ModelError(f"references[{index}] has missing or unexpected fields")
        converted = {field: _text(item[field], f"references[{index}].{field}") for field in fields if field != "url"}
        converted["url"] = _public_https(item["url"], f"references[{index}].url")
        if not REFERENCE_ID.fullmatch(converted["id"]) or converted["id"] in seen_references:
            raise ModelError(f"invalid or duplicate reference id: {converted['id']}")
        seen_references.add(converted["id"])
        references.append(converted)
    model["references"] = references

    glossary: list[dict[str, str]] = []
    for index, item in enumerate(_list(raw["glossary"], "glossary", minimum=5, maximum=100)):
        if not isinstance(item, dict) or set(item) != {"term", "definition"}:
            raise ModelError(f"glossary[{index}] has missing or unexpected fields")
        glossary.append({"term": _text(item["term"], f"glossary[{index}].term"), "definition": _text(item["definition"], f"glossary[{index}].definition")})
    model["glossary"] = glossary

    figure = raw["figure"]
    if not isinstance(figure, dict) or set(figure) != {"title", "alt_text", "nodes", "edges"}:
        raise ModelError("figure has missing or unexpected fields")
    nodes = [_text(item, "figure.nodes") for item in _list(figure["nodes"], "figure.nodes", minimum=3, maximum=7)]
    edges: list[list[str]] = []
    for index, edge in enumerate(_list(figure["edges"], "figure.edges", minimum=2, maximum=12)):
        values = [_text(item, f"figure.edges[{index}]") for item in _list(edge, f"figure.edges[{index}]", minimum=3, maximum=3)]
        if values[0] not in nodes or values[1] not in nodes:
            raise ModelError(f"figure edge references an unknown node: {values}")
        edges.append(values)
    model["figure"] = {"title": _text(figure["title"], "figure.title"), "alt_text": _text(figure["alt_text"], "figure.alt_text"), "nodes": nodes, "edges": edges}

    derivative = raw["derivative_sources"]
    derivative_fields = {"cover", "executive_outcomes", "approval_conditions", "figure", "findings", "decision", "review_trigger"}
    if not isinstance(derivative, dict) or set(derivative) != derivative_fields:
        raise ModelError("derivative_sources has missing or unexpected fields")
    known_sources = seen_evidence | seen_references

    def source_ids(value: object, field: str) -> list[str]:
        ids = [_text(item, field) for item in _list(value, field, minimum=1, maximum=20)]
        if len(ids) != len(set(ids)) or any(item not in known_sources for item in ids):
            raise ModelError(f"{field} contains duplicate or unknown source ids")
        return ids

    outcome_sources = _list(derivative["executive_outcomes"], "derivative_sources.executive_outcomes", minimum=len(model["executive_outcomes"]), maximum=len(model["executive_outcomes"]))
    condition_sources = _list(derivative["approval_conditions"], "derivative_sources.approval_conditions", minimum=len(model["approval_conditions"]), maximum=len(model["approval_conditions"]))
    finding_sources = derivative["findings"]
    if not isinstance(finding_sources, dict) or set(finding_sources) != seen_findings:
        raise ModelError("derivative_sources.findings must cover every finding exactly")
    model["derivative_sources"] = {
        "cover": source_ids(derivative["cover"], "derivative_sources.cover"),
        "executive_outcomes": [source_ids(item, f"derivative_sources.executive_outcomes[{index}]") for index, item in enumerate(outcome_sources)],
        "approval_conditions": [source_ids(item, f"derivative_sources.approval_conditions[{index}]") for index, item in enumerate(condition_sources)],
        "figure": source_ids(derivative["figure"], "derivative_sources.figure"),
        "findings": {finding_id: source_ids(finding_sources[finding_id], f"derivative_sources.findings.{finding_id}") for finding_id in sorted(seen_findings)},
        "decision": source_ids(derivative["decision"], "derivative_sources.decision"),
        "review_trigger": source_ids(derivative["review_trigger"], "derivative_sources.review_trigger"),
    }
    return model


def _assert_evidence_bindings(model: dict[str, object], claim_evidence: list[tuple[str, tuple[str, ...]]]) -> None:
    known = {str(item["id"]) for item in model["evidence"]}
    for claim_id, evidence_ids in claim_evidence:
        missing = sorted(set(evidence_ids) - known)
        if missing:
            raise ModelError(f"claim {claim_id} references unknown evidence ids: {missing}")
    for finding in model["findings"]:
        ids = [item.strip() for item in re.split(r"[,;]", str(finding["evidence_ids"])) if item.strip()]
        if not ids or any(not EVIDENCE_ID.fullmatch(item) for item in ids):
            raise ModelError(f"finding {finding['id']} has invalid evidence ids")
        missing = sorted(set(ids) - known)
        if missing:
            raise ModelError(f"finding {finding['id']} references unknown evidence ids: {missing}")


def load_bound_report_model(stage_root: Path) -> tuple[dict[str, object], list[dict[str, object]], str, str, dict[str, str]]:
    bundle = consume_validated_stages(stage_root)
    _, manifest_bytes = read_bounded_regular_file(stage_root, "stage-manifest.json", 1024 * 1024)
    manifest = json.loads(manifest_bytes.decode("utf-8"))
    entry = manifest.get("report_model")
    if not isinstance(entry, dict) or entry.get("status") != "Validated":
        raise ModelError("stage manifest lacks a validated report_model entry")
    filename = entry.get("file")
    if not isinstance(filename, str) or not filename:
        raise ModelError("report_model filename is invalid")
    path, data = read_bounded_regular_file(stage_root, filename, MAX_MODEL_BYTES)
    digest = hashlib.sha256(data).hexdigest()
    if entry.get("sha256") != digest:
        raise ModelError(f"report model digest mismatch: {path}")
    try:
        raw = json.loads(data.decode("utf-8", errors="strict"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ModelError("report model is invalid JSON") from exc
    model = validate_report_model(raw)
    if model["assessment"] != bundle.assessment or model["target"] != bundle.target or model["version"] != bundle.version:
        raise ModelError("report model identity does not match validated stage claims")
    _assert_evidence_bindings(model, [(claim_id, claim.evidence_ids) for claim_id, claim in bundle.claims.items()])
    claims_entry = manifest.get("claims")
    if not isinstance(claims_entry, dict) or not isinstance(claims_entry.get("file"), str) or not isinstance(claims_entry.get("sha256"), str):
        raise ModelError("stage manifest claims binding is invalid")
    return model, manifest["stages"], digest, hashlib.sha256(manifest_bytes).hexdigest(), {"file": claims_entry["file"], "sha256": claims_entry["sha256"]}


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _column_widths(columns: list[str], rows: list[list[str]]) -> list[int]:
    weights = []
    for index, heading in enumerate(columns):
        longest = max([len(heading)] + [min(len(row[index]), 240) for row in rows])
        weights.append(max(8, min(longest, 80)))
    total = sum(weights)
    raw = [max(720, round(9360 * weight / total)) for weight in weights]
    scale = 9360 / sum(raw)
    result = [round(value * scale) for value in raw]
    result[-1] += 9360 - sum(result)
    return result


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _bookmark(paragraph, name: str, text: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    run = paragraph.add_run(text)
    run.bold = True
    paragraph._p.insert(paragraph._p.index(run._r), start)
    paragraph._p.insert(paragraph._p.index(run._r) + 1, end)


def _internal_link(paragraph, label: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _external_link(paragraph, label: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    properties.append(style)
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _complex_field(paragraph, instruction: str) -> None:
    begin_run = OxmlElement("w:r")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run.append(begin)
    instruction_run = OxmlElement("w:r")
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_text.text = f" {instruction} "
    instruction_run.append(instruction_text)
    separate_run = OxmlElement("w:r")
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run.append(separate)
    result_run = OxmlElement("w:r")
    result_text = OxmlElement("w:t")
    result_text.text = "1"
    result_run.append(result_text)
    end_run = OxmlElement("w:r")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run.append(end)
    for element in (begin_run, instruction_run, separate_run, result_run, end_run):
        paragraph._p.append(element)


def _add_cited_text(paragraph, value: str) -> None:
    cursor = 0
    for match in CITATION.finditer(value):
        if match.start() > cursor:
            paragraph.add_run(value[cursor:match.start()])
        label = match.group(0)
        _internal_link(paragraph, label, f"REF_{match.group(1)}")
        cursor = match.end()
    if cursor < len(value):
        paragraph.add_run(value[cursor:])


def _add_table(doc: Document, title: str, columns: list[str], rows: list[list[str]], *, compact: bool = False) -> None:
    caption = doc.add_paragraph(style="Caption")
    caption.add_run(title).bold = True
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header = table.rows[0]
    _repeat_header(header)
    for index, value in enumerate(columns):
        cell = header.cells[index]
        cell.text = value
        _shade_cell(cell, "F2F4F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            _add_cited_text(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(8)
    _set_table_geometry(table, _column_widths(columns, rows))
    if compact:
        for row in table.rows:
            for cell in row.cells:
                _set_cell_margins(cell, top=35, bottom=35)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def _configure_document(doc: Document, model: dict[str, object]) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = f"{model['target']} | Security assessment | Classification: PUBLIC"
    header.style = styles["Header"]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("PUBLIC | Page ")
    _complex_field(footer, "PAGE")

    properties = doc.core_properties
    properties.author = "Security Assessment Automation"
    properties.last_modified_by = "Security Assessment Automation"
    properties.title = f"{model['target']} security assessment"
    properties.subject = "Evidence-led Microsoft IDE extension security assessment"
    properties.keywords = "security assessment, VS Code, AI, OWASP, privacy"
    properties.comments = "Generated from hash-bound validated stage outputs."


def _draw_figure(
    model: dict[str, object],
    output: Path,
    output_identity: FileIdentity,
) -> None:
    figure = model["figure"]
    assert isinstance(figure, dict)
    nodes = figure["nodes"]
    edges = figure["edges"]
    assert isinstance(nodes, list) and isinstance(edges, list)
    canvas = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(canvas)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 40)
        body_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 28)
        small_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 23)
    except OSError:
        title_font = body_font = small_font = ImageFont.load_default()
    draw.text((70, 45), str(figure["title"]), fill="#0B2545", font=title_font)
    columns = max(2, min(4, len(nodes)))
    box_width, box_height = 330, 150
    positions: dict[str, tuple[int, int, int, int]] = {}
    for index, node in enumerate(nodes):
        row, column = divmod(index, columns)
        x = 80 + column * 420
        y = 180 + row * 300
        positions[str(node)] = (x, y, x + box_width, y + box_height)
    for source, target, label in edges:
        source_box = positions[str(source)]
        target_box = positions[str(target)]
        start = ((source_box[0] + source_box[2]) // 2, (source_box[1] + source_box[3]) // 2)
        end = ((target_box[0] + target_box[2]) // 2, (target_box[1] + target_box[3]) // 2)
        draw.line((start, end), fill="#49677F", width=5)
        midpoint = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2)
        draw.text((midpoint[0] + 8, midpoint[1] - 30), str(label)[:34], fill="#49677F", font=small_font)
    for node, box in positions.items():
        draw.rounded_rectangle(box, radius=18, fill="#E8EEF5", outline="#2E74B5", width=4)
        words = node.split()
        lines: list[str] = []
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if len(candidate) > 22 and current:
                lines.append(current)
                current = word
            else:
                current = candidate
        if current:
            lines.append(current)
        y = box[1] + 28
        for line in lines[:3]:
            bbox = draw.textbbox((0, 0), line, font=body_font)
            draw.text((box[0] + (box_width - (bbox[2] - bbox[0])) // 2, y), line, fill="#0B2545", font=body_font)
            y += 38
    with _open_owned_file(
        output,
        output_identity,
        "architecture figure",
        write=True,
    ) as stream:
        canvas.save(stream, format="PNG", optimize=True)
    _require_regular_identity(output, output_identity, "architecture figure")


def _assert_contents_starts_on_fresh_page(doc: Document) -> None:
    contents = [
        index
        for index, paragraph in enumerate(doc.paragraphs)
        if paragraph.text == "Contents"
    ]
    if len(contents) != 1:
        raise ModelError("report must contain exactly one Contents heading")
    index = contents[0]
    if index == 0:
        raise ModelError("Contents must be preceded by an explicit page break")
    page_breaks = [
        node
        for node in doc.paragraphs[index - 1]._p.iter(qn("w:br"))
        if node.get(qn("w:type")) == "page"
    ]
    if not page_breaks:
        raise ModelError("Contents must start on a fresh page")


def _build_report_owned(
    model: dict[str, object],
    output: Path,
    figure_path: Path,
    output_identity: FileIdentity,
    figure_identity: FileIdentity,
) -> None:
    doc = Document()
    _configure_document(doc, model)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(model["target"]))
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Microsoft IDE AI Extension Security Assessment").italic = True
    classification = doc.add_paragraph()
    classification.alignment = WD_ALIGN_PARAGRAPH.CENTER
    classification.add_run("Classification: PUBLIC").bold = True

    doc.add_heading("Document Control", level=1)
    _add_table(doc, "Table 1 — Document control", ["Field", "Value"], [
        ["Assessment", str(model["assessment"])],
        ["Target", str(model["target"])],
        ["Publisher", str(model["publisher"])],
        ["Extension ID", str(model["extension_id"])],
        ["Version", str(model["version"])],
        ["Run key", str(model["run_key"])],
        ["Classification", "PUBLIC"],
        ["Distribution", "Public distribution"],
        ["Decision", str(model["decision"])],
    ])
    doc.add_heading("Revision History", level=1)
    _add_table(
        doc,
        "Table 2 — Revision history",
        ["Document version", "Date", "Status", "Change"],
        [
            [
                str(item["version"]),
                str(item["date"]),
                str(item["status"]),
                str(item["change"]),
            ]
            for item in model["revision_history"]
        ],
    )
    # Contents is a navigation surface, not trailing front matter. A mandatory
    # page break prevents it from ever beginning in the lower half of a page.
    doc.add_page_break()
    doc.add_heading("Contents", level=1)
    toc = doc.add_paragraph()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    toc._p.append(field)
    for section in model["sections"]:
        assert isinstance(section, dict)
        item = doc.add_paragraph(style="List Bullet")
        item.paragraph_format.space_before = Pt(0)
        item.paragraph_format.space_after = Pt(0)
        item.paragraph_format.line_spacing = 1.0
        item.add_run(str(section["heading"])).font.size = Pt(9)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Executive Summary", level=1)
    lead = doc.add_paragraph()
    _add_cited_text(lead, f"Decision: {model['decision']}. Overall residual risk: {model['overall_residual_risk']}. This assessment is scoped to {', '.join(model['ide_scope'])} version {model['version']} and is based on the evidence state and limitations recorded below.")
    for outcome in model["executive_outcomes"]:
        paragraph = doc.add_paragraph(style="List Bullet")
        _add_cited_text(paragraph, str(outcome))
    doc.add_heading("Approval Conditions and Immediate Actions", level=2)
    for condition in model["approval_conditions"]:
        paragraph = doc.add_paragraph(style="List Number")
        _add_cited_text(paragraph, str(condition))

    doc.add_heading("Architecture and Trust-Boundary Figure", level=2)
    _draw_figure(model, figure_path, figure_identity)
    paragraph = doc.add_paragraph()
    with _open_owned_file(
        figure_path,
        figure_identity,
        "architecture figure",
        write=False,
    ) as stream:
        inline = paragraph.add_run().add_picture(stream, width=Inches(6.5))._inline
    inline.docPr.set("name", str(model["figure"]["title"]))
    inline.docPr.set("title", str(model["figure"]["title"]))
    inline.docPr.set("descr", str(model["figure"]["alt_text"]))
    caption = doc.add_paragraph(style="Caption")
    caption.add_run(f"Figure 1 — {model['figure']['title']}")

    for section in model["sections"]:
        assert isinstance(section, dict)
        doc.add_heading(str(section["heading"]), level=int(section["level"]))
        if str(section["heading"]).casefold().startswith("part ii") and "Visual Studio" not in model["ide_scope"]:
            doc.add_paragraph(
                "Decision: Not applicable. Visual Studio is unsupported and out of scope for this VS Code extension. "
                "Evidence: the Marketplace manifest declares the Visual Studio Code engine only. Limitation: no Visual Studio package or runtime was assessed. "
                "Verification: reassess if the publisher releases a Visual Studio extension or claims Visual Studio compatibility."
            )
        for value in section["paragraphs"]:
            paragraph = doc.add_paragraph()
            _add_cited_text(paragraph, str(value))
        for value in section["bullets"]:
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_cited_text(paragraph, str(value))
        for table in section["tables"]:
            assert isinstance(table, dict)
            _add_table(doc, str(table["title"]), list(table["columns"]), list(table["rows"]))

    doc.add_heading("Individual Finding Records", level=1)
    for finding in model["findings"]:
        assert isinstance(finding, dict)
        doc.add_heading(f"{finding['id']} — {finding['title']}", level=2)
        _add_table(doc, f"Finding record — {finding['id']}", ["Field", "Assessment"], [[field.replace("_", " ").title(), str(finding[field])] for field in (
            "scope", "scenario", "evidence_ids", "likelihood", "impact", "inherent", "controls", "control_strength", "residual_likelihood", "residual_impact", "residual", "recommendation", "owner", "priority", "target_date", "verification", "mappings", "confidence"
        )])

    doc.add_heading("Consolidated Risk Register", level=1)
    doc.add_paragraph("Legend: ID is the finding identifier; L and I are inherent likelihood and impact scored 1–5; Inherent is L × I; rL and rI are residual likelihood and impact after verified controls; Residual is rL × rI; rating bands are 1–4 Low, 5–9 Moderate, 10–16 High, and 17–25 Critical. Treatment is the required action and Owner is accountable for verification.")
    _add_table(doc, "Consolidated risk register", ["ID", "Risk", "Scope", "L/I", "Inherent", "rL/rI", "Residual", "Treatment", "Owner"], [[str(item["id"]), str(item["title"]), str(item["scope"]), f"{item['likelihood']}/{item['impact']}", str(item["inherent"]), f"{item['residual_likelihood']}/{item['residual_impact']}", str(item["residual"]), str(item["recommendation"]), str(item["owner"])] for item in model["findings"]])

    doc.add_heading("Evidence Register", level=1)
    _add_table(doc, "Evidence register", ["ID", "Evidence", "Source", "Method", "State", "Limitation"], [[str(item[field]) for field in ("id", "title", "source", "method", "state", "limitation")] for item in model["evidence"]])

    doc.add_heading("References", level=1)
    reference_table = doc.add_table(rows=1, cols=6)
    reference_table.style = "Table Grid"
    for index, heading in enumerate(("ID", "Title", "Publisher", "Source", "Accessed", "Applicability")):
        reference_table.rows[0].cells[index].text = heading
        _shade_cell(reference_table.rows[0].cells[index], "F2F4F7")
    _repeat_header(reference_table.rows[0])
    bookmark_id = 1000
    reference_rows: list[list[str]] = []
    for item in model["references"]:
        assert isinstance(item, dict)
        row = reference_table.add_row()
        _bookmark(row.cells[0].paragraphs[0], str(item["id"]).replace("-", "_"), str(item["id"]), bookmark_id)
        bookmark_id += 1
        row.cells[1].text = str(item["title"])
        row.cells[2].text = str(item["publisher"])
        _external_link(row.cells[3].paragraphs[0], "Open source", str(item["url"]))
        row.cells[4].text = str(item["accessed"])
        row.cells[5].text = str(item["applicability"])
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
        reference_rows.append([str(item[field]) for field in ("id", "title", "publisher", "url", "accessed", "applicability")])
    _set_table_geometry(reference_table, _column_widths(["ID", "Title", "Publisher", "Source", "Accessed", "Applicability"], reference_rows))

    doc.add_heading("Appendices", level=1)
    doc.add_heading("Appendix A — Analysis Selections and Results", level=2)
    doc.add_paragraph("Static analysis: selected and completed within the recorded package scope. Static malware review: selected and completed within the available scanner scope. Runtime analysis: selected; any unavailable representative host is recorded as Blocked rather than silently treated as No. Privacy review: selected and completed from package, policy, documentation, and observed evidence. Evidence dispositions used by this report are Verified, Inferred, Not observed, Not applicable, Unknown, and Blocked. Native Microsoft Word closeout remains a separate final gate.")
    doc.add_heading("Appendix B — Build and Quality-Assurance Contract", level=2)
    doc.add_paragraph("The candidate report was compiled from all fifteen hash-bound validated stage outputs and a hash-bound report model. Deterministic validation, publication-safety inspection, accessibility review, privacy and metadata review, citation and bookmark audit, scoring reconciliation, secure review, native Microsoft Word rendering, page count, and every-page inspection are shipping gates. A passed candidate build alone is not an authoritative determination.")
    doc.add_heading("Glossary", level=1)
    _add_table(doc, "Glossary", ["Term", "Definition"], [[str(item["term"]), str(item["definition"])] for item in model["glossary"]], compact=True)

    _assert_contents_starts_on_fresh_page(doc)
    with _open_owned_file(output, output_identity, "output", write=True) as stream:
        doc.save(stream)
    _require_regular_identity(output, output_identity, "output")


def build_report(
    model: dict[str, object],
    output: Path,
    figure_path: Path,
    *,
    output_identity: FileIdentity | None = None,
    figure_identity: FileIdentity | None = None,
) -> None:
    created: list[tuple[Path, FileIdentity]] = []
    try:
        if figure_identity is None:
            figure_identity = _create_owned_file(figure_path, "architecture figure")
            created.append((figure_path, figure_identity))
        if output_identity is None:
            output_identity = _create_owned_file(output, "output")
            created.append((output, output_identity))
        _build_report_owned(
            model,
            output,
            figure_path,
            output_identity,
            figure_identity,
        )
    except Exception:
        for path, expected in reversed(created):
            _unlink_if_owned(path, expected)
        raise


def write_build_manifest(
    model: dict[str, object],
    stages: list[dict[str, object]],
    stage_root: Path,
    output: Path,
    manifest_path: Path,
    report_model_sha256: str,
    stage_manifest_sha256: str,
    claims: dict[str, str],
    *,
    workspace_root: Path,
    output_source: Path | None = None,
    output_source_identity: FileIdentity | None = None,
    manifest_identity: FileIdentity | None = None,
) -> None:
    stage_hashes = {str(item["file"]): str(item["sha256"]) for item in stages}
    digest_source = output if output_source is None else output_source
    if output_source_identity is None:
        try:
            source_metadata = digest_source.lstat()
        except FileNotFoundError as exc:
            raise ModelError("output disappeared before manifest generation") from exc
        if is_link_or_reparse(source_metadata) or not stat.S_ISREG(
            source_metadata.st_mode
        ):
            raise ModelError("output must be a regular non-symlink file")
        output_source_identity = _identity(source_metadata)
    output_sha256 = _sha256_owned_file(
        digest_source,
        output_source_identity,
        "output",
    )
    payload = {
        "schema_version": 1,
        "assessment": model["assessment"],
        "run_key": model["run_key"],
        "parent_skill_version": "1.4.5",
        "design_preset": "standard_business_brief",
        "header_pattern": "memo_masthead",
        "generated_at": datetime.now(UTC).isoformat(),
        "python": sys.version.split()[0],
        "python_docx": importlib.metadata.version("python-docx"),
        "pillow": importlib.metadata.version("Pillow"),
        "stage_root": str(stage_root.relative_to(workspace_root)),
        "stage_sha256": stage_hashes,
        "stage_manifest_sha256": stage_manifest_sha256,
        "claims": claims,
        "report_model_sha256": report_model_sha256,
        "output": str(output.relative_to(workspace_root)),
        "output_sha256": output_sha256,
        "native_word_closeout": "Pending",
    }
    created = False
    try:
        if manifest_identity is None:
            manifest_identity = _create_owned_file(manifest_path, "build manifest")
            created = True
        data = (json.dumps(payload, indent=2, sort_keys=True) + "\n").encode("utf-8")
        with _open_owned_file(
            manifest_path,
            manifest_identity,
            "build manifest",
            write=True,
        ) as stream:
            stream.write(data)
        _require_regular_identity(manifest_path, manifest_identity, "build manifest")
    except Exception:
        if created:
            _unlink_if_owned(manifest_path, manifest_identity)
        raise


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--workspace-root", required=True, type=Path)
    parser.add_argument("--stage-root", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--build-manifest", required=True, type=Path)
    args = parser.parse_args(argv)
    workspace_root = _resolve_workspace_root(args.workspace_root)
    stage_root = _resolve_stage_root(args.stage_root, workspace_root)
    output = _resolve_destination(args.output, "output", workspace_root)
    manifest_path = _resolve_destination(
        args.build_manifest, "build manifest", workspace_root
    )
    model, stages, report_model_sha256, stage_manifest_sha256, claims = load_bound_report_model(stage_root)
    figure_path = _resolve_destination(
        manifest_path.parent / f"{model['run_key']}-architecture.png",
        "architecture figure",
        workspace_root,
    )
    if len({output, manifest_path, figure_path}) != 3:
        raise ModelError("output, architecture figure, and build manifest must be distinct")

    staged: list[tuple[Path, Path, FileIdentity, str]] = []
    try:
        for destination, label in (
            (figure_path, "architecture figure"),
            (output, "output"),
            (manifest_path, "build manifest"),
        ):
            temporary, identity = _create_secure_temp(destination, label)
            staged.append((temporary, destination, identity, label))

        figure_temp, _, figure_identity, _ = staged[0]
        output_temp, _, output_identity, _ = staged[1]
        manifest_temp, _, manifest_identity, _ = staged[2]
        build_report(
            model,
            output_temp,
            figure_temp,
            output_identity=output_identity,
            figure_identity=figure_identity,
        )
        write_build_manifest(
            model,
            stages,
            stage_root,
            output,
            manifest_temp,
            report_model_sha256,
            stage_manifest_sha256,
            claims,
            workspace_root=workspace_root,
            output_source=output_temp,
            output_source_identity=output_identity,
            manifest_identity=manifest_identity,
        )
        _publish_bundle(staged)
    finally:
        for temporary, _, expected, _ in reversed(staged):
            _unlink_if_owned(temporary, expected)
    print(f"Built {output}")
    print(f"Build manifest {manifest_path}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (ModelError, ValueError, OSError) as exc:
        print(f"build_assessment_docx: {exc}", file=sys.stderr)
        raise SystemExit(2) from None
