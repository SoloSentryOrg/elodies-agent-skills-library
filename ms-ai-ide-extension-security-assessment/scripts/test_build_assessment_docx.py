from __future__ import annotations

import importlib.util
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

SCRIPTS = Path(__file__).parent
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))
SPEC = importlib.util.spec_from_file_location(
    "build_assessment_docx", SCRIPTS / "build_assessment_docx.py"
)
assert SPEC and SPEC.loader
MODULE = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = MODULE
SPEC.loader.exec_module(MODULE)


def _model() -> dict[str, object]:
    long_paragraph = (
        "Verified package evidence shows a bounded security-relevant behavior and its "
        "enterprise consequence. The assessment distinguishes directly observed package "
        "facts from inference, records the affected trust boundary, and links the claim to "
        "REF-001 and REF-002. The control recommendation applies least privilege, explicit "
        "approval, secure defaults, monitoring, and a repeatable verification test. "
    ) * 4
    required_headings = [
        "Purpose & Function Overview",
        "Scope, Assumptions, Methodology and Assessment Criteria",
        "Product Identity and Version",
        "Part I — VS Code Assessment",
        "VS Code Executive Decision and Approval",
        "VS Code Purpose and Function",
        "VS Code Architecture",
        "VS Code Trust Boundaries",
        "VS Code Data Flows",
        "VS Code Installation Manifest and Uninstall Behavior",
        "VS Code Runtime Processes and Activation",
        "VS Code Network Communications, Authentication, Telemetry and Privacy",
        "VS Code MCP and Agent Skills",
        "VS Code Supply Chain and Dependencies",
        "VS Code Threats and Findings",
        "VS Code Framework Disposition",
        "VS Code Controls and Detection",
        "VS Code Residual Risk, Confidence, Limitations, Evidence and Verification",
        "Part II — Visual Studio Assessment",
        "Visual Studio Decision, Evidence, Limitations and Verification",
        "Part III — Installed Agent Skills",
        "Cross-IDE Comparison",
        "Consolidated Supply Chain",
        "Consolidated Privacy and Data Protection Assessment",
        "Enterprise Controls Roadmap",
        "Detection Opportunities and Monitoring Plan",
        "Residual Risk and Approval Recommendation",
        "Limitations and Confidence",
    ]
    sections = []
    for index, heading in enumerate(required_headings):
        tables = []
        if index < 12:
            tables.append({
                "title": f"Table {index + 3} — {heading}",
                "columns": ["Control", "Evidence", "Disposition"],
                "rows": [["Least privilege", "EVD-001 and REF-001", "Required"], ["Runtime", "EVD-002", "Blocked pending representative host"]],
            })
        sections.append({
            "id": f"section-{index + 1}",
            "heading": heading,
            "level": 1 if heading.startswith(("Part ", "Cross-IDE", "Consolidated", "Privacy", "Enterprise", "Detection", "Residual", "Limitations", "Purpose", "Scope", "Product")) else 2,
            "paragraphs": [long_paragraph],
            "bullets": ["Verified control statement supported by EVD-001 and REF-003."],
            "tables": tables,
        })
    finding_fields = {
        "id": "F-001",
        "title": "Representative high-agency extension risk",
        "scope": "VS Code 1.100 or later",
        "scenario": "Untrusted workspace content influences an agent that can edit files and invoke commands. Evidence EVD-001 supports the package capability; runtime confirmation is blocked.",
        "evidence_ids": "EVD-001; EVD-002",
        "likelihood": "4 Likely",
        "impact": "4 Major",
        "inherent": "16 High",
        "controls": "Workspace Trust, explicit approvals, restricted egress, protected branches, and synthetic-data testing.",
        "control_strength": "1 Weak because controls are proposed rather than observed in a representative host.",
        "residual_likelihood": "3 Possible",
        "residual_impact": "4 Major",
        "residual": "12 High",
        "recommendation": "Use a managed disposable profile and prohibit production credentials until runtime verification passes.",
        "owner": "Endpoint Security Owner",
        "priority": "P1",
        "target_date": "Before deployment",
        "verification": "Exercise representative workflows with restricted egress and confirm approval, file, command and cleanup evidence.",
        "mappings": "OWASP LLM prompt injection and excessive agency; NIST AI RMF MEASURE and MANAGE.",
        "confidence": "Medium",
    }
    references = [
        {
            "id": f"REF-{index:03d}",
            "title": f"Primary source {index}",
            "publisher": "Example Publisher",
            "url": f"https://example.com/source-{index}",
            "accessed": "2026-08-01",
            "applicability": "Version-specific primary evidence",
        }
        for index in range(1, 13)
    ]
    return {
        "schema_version": 2,
        "assessment": "Example MCP",
        "target": "Example MCP",
        "publisher": "Example Publisher",
        "extension_id": "example.publisher",
        "version": "1.2.3",
        "run_key": "2026-08-01-v1.0",
        "assessment_date": "2026-08-01",
        "document_version": "1.0",
        "classification": "PUBLIC",
        "decision": "Defer pending evidence",
        "overall_residual_risk": "High",
        "review_trigger": "Representative runtime evidence or version change",
        "ide_scope": ["VS Code"],
        "executive_outcomes": [f"Outcome {index} is supported by REF-{index:03d} and retained package evidence." for index in range(1, 6)],
        "approval_conditions": ["Complete representative isolated runtime testing before production approval."],
        "sections": sections,
        "findings": [finding_fields],
        "evidence": [
            {"id": f"EVD-{index:03d}", "title": f"Evidence {index}", "source": "Synthetic test fixture", "method": "Static inspection", "state": "Verified", "limitation": "No representative runtime"}
            for index in range(1, 6)
        ],
        "references": references,
        "glossary": [
            {"term": "Agent", "definition": "A model-directed component that may invoke tools."},
            {"term": "MCP", "definition": "Model Context Protocol."},
            {"term": "Residual risk", "definition": "Risk remaining after verified controls."},
            {"term": "VSIX", "definition": "Visual Studio Code extension package."},
            {"term": "Workspace Trust", "definition": "VS Code control for untrusted workspaces."},
        ],
        "figure": {
            "title": "Example trust boundaries",
            "alt_text": "A diagram connecting an untrusted workspace to the extension host, local tools, and external services.",
            "nodes": ["Untrusted workspace", "VS Code extension host", "Local tools", "External services"],
            "edges": [["Untrusted workspace", "VS Code extension host", "files and prompts"], ["VS Code extension host", "Local tools", "commands"], ["VS Code extension host", "External services", "HTTPS"]],
        },
        "derivative_sources": {
            "cover": ["EVD-001", "REF-001"],
            "executive_outcomes": [[f"REF-{index:03d}", "EVD-001"] for index in range(1, 6)],
            "approval_conditions": [["EVD-001", "REF-001"]],
            "figure": ["EVD-001", "REF-001"],
            "findings": {"F-001": ["EVD-001", "EVD-002"]},
            "decision": ["EVD-001"],
            "review_trigger": ["EVD-001", "REF-001"],
        },
    }


class AssessmentDocxBuilderTests(unittest.TestCase):
    def test_revision_history_accepts_correction_and_requires_latest_version(self) -> None:
        raw = _model()
        raw["document_version"] = "1.1"
        raw["revision_history"] = [
            {
                "version": "1.0",
                "date": "2026-08-01",
                "status": "Assessment complete; deferred",
                "change": "Initial evidence-led assessment.",
            },
            {
                "version": "1.1",
                "date": "2026-08-01",
                "status": "Citation correction; deferred",
                "change": "Corrected one external reference; findings unchanged.",
            },
        ]
        model = MODULE.validate_report_model(raw)
        self.assertEqual(len(model["revision_history"]), 2)
        self.assertEqual(model["revision_history"][-1]["version"], "1.1")

        raw["revision_history"][-1]["version"] = "1.2"
        with self.assertRaisesRegex(MODULE.ModelError, "must match"):
            MODULE.validate_report_model(raw)

        raw["revision_history"][-1]["version"] = "1.1"
        raw["revision_history"][-1]["change"] = "x" * 4097
        with self.assertRaisesRegex(MODULE.ModelError, "exceeds 4096 bytes"):
            MODULE.validate_report_model(raw)

    def test_revision_history_rejects_duplicates_invalid_dates_and_order(self) -> None:
        raw = _model()
        raw["document_version"] = "1.1"
        raw["revision_history"] = [
            {"version": "1.0", "date": "2026-08-02", "status": "Complete", "change": "Initial."},
            {"version": "1.1", "date": "2026-08-01", "status": "Corrected", "change": "Citation."},
        ]
        with self.assertRaisesRegex(MODULE.ModelError, "nondecreasing"):
            MODULE.validate_report_model(raw)

        raw["revision_history"][0]["date"] = "2026-02-30"
        raw["revision_history"][1]["date"] = "2026-08-01"
        with self.assertRaisesRegex(MODULE.ModelError, "valid ISO date"):
            MODULE.validate_report_model(raw)

        raw["revision_history"][0]["date"] = "2026-08-01"
        raw["revision_history"][1]["version"] = "1.0"
        with self.assertRaisesRegex(MODULE.ModelError, "invalid or duplicate"):
            MODULE.validate_report_model(raw)

        raw["revision_history"][1]["version"] = "v1.1"
        with self.assertRaisesRegex(MODULE.ModelError, "invalid or duplicate"):
            MODULE.validate_report_model(raw)

        raw["revision_history"][0]["version"] = "2.0"
        raw["revision_history"][1]["version"] = "1.1"
        with self.assertRaisesRegex(MODULE.ModelError, "in increasing order"):
            MODULE.validate_report_model(raw)

    def test_destination_rejects_dangling_leaf_symlink_without_following_it(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory).resolve()
            parent = root / "reports"
            parent.mkdir()
            destination = parent / "assessment.docx"
            destination.symlink_to("missing.docx")
            with self.assertRaisesRegex(MODULE.ModelError, "must not be a symlink"):
                MODULE._resolve_destination(destination, "output", root)
            self.assertTrue(destination.is_symlink())

    def test_build_report_refuses_to_overwrite_and_cleans_only_its_figure(self) -> None:
        model = MODULE.validate_report_model(_model())
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            output = root / "assessment.docx"
            output.write_bytes(b"preexisting report")
            figure = root / "figure.png"
            with self.assertRaisesRegex(MODULE.ModelError, "already exists"):
                MODULE.build_report(model, output, figure)
            self.assertEqual(output.read_bytes(), b"preexisting report")
            self.assertFalse(figure.exists())

    def test_exclusive_publication_preserves_preexisting_target(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            temporary = root / "temporary.docx"
            expected = MODULE._create_owned_file(temporary, "temporary output")
            destination = root / "assessment.docx"
            destination.write_bytes(b"preexisting report")
            with self.assertRaisesRegex(MODULE.ModelError, "refusing to overwrite"):
                MODULE._publish_exclusive(
                    temporary,
                    destination,
                    expected,
                    "output",
                )
            self.assertEqual(destination.read_bytes(), b"preexisting report")

    def test_build_manifest_refuses_to_overwrite_existing_file(self) -> None:
        model = MODULE.validate_report_model(_model())
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory).resolve()
            stage_root = root / "stage-output"
            reports = root / "reports"
            stage_root.mkdir()
            reports.mkdir()
            output = reports / "assessment.docx"
            output.write_bytes(b"candidate report")
            manifest = reports / "assessment.build.json"
            manifest.write_bytes(b"preexisting manifest")
            with self.assertRaisesRegex(MODULE.ModelError, "already exists"):
                MODULE.write_build_manifest(
                    model,
                    [],
                    stage_root,
                    output,
                    manifest,
                    "a" * 64,
                    "b" * 64,
                    {"file": "claims-register.json", "sha256": "c" * 64},
                    workspace_root=root,
                )
            self.assertEqual(manifest.read_bytes(), b"preexisting manifest")

    def test_destination_rejects_path_outside_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory).resolve()
            workspace = root / "workspace"
            outside = root / "outside"
            workspace.mkdir()
            outside.mkdir()
            with self.assertRaisesRegex(MODULE.ModelError, "workspace root"):
                MODULE._resolve_destination(outside / "report.docx", "output", workspace)

    def test_publication_rollback_does_not_delete_raced_replacements(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            first_temp = root / "first.tmp"
            second_temp = root / "second.tmp"
            first_identity = MODULE._create_owned_file(first_temp, "first temporary")
            second_identity = MODULE._create_owned_file(second_temp, "second temporary")
            first_destination = root / "first.docx"
            second_destination = root / "second.json"
            real_link = os.link
            calls = 0

            def racing_link(
                source: os.PathLike[str],
                destination: os.PathLike[str],
                *,
                follow_symlinks: bool = True,
            ) -> None:
                nonlocal calls
                calls += 1
                if calls == 1:
                    real_link(
                        source,
                        destination,
                        follow_symlinks=follow_symlinks,
                    )
                    return
                first_destination.unlink()
                first_destination.write_bytes(b"raced first replacement")
                Path(destination).write_bytes(b"raced second replacement")
                raise FileExistsError

            with mock.patch.object(MODULE.os, "link", side_effect=racing_link):
                with self.assertRaisesRegex(MODULE.ModelError, "appeared during generation"):
                    MODULE._publish_bundle(
                        [
                            (
                                first_temp,
                                first_destination,
                                first_identity,
                                "output",
                            ),
                            (
                                second_temp,
                                second_destination,
                                second_identity,
                                "build manifest",
                            ),
                        ]
                    )
            self.assertEqual(first_destination.read_bytes(), b"raced first replacement")
            self.assertEqual(second_destination.read_bytes(), b"raced second replacement")

    def test_contents_invariant_rejects_heading_without_fresh_page(self) -> None:
        from docx import Document

        document = Document()
        document.add_paragraph("Front matter")
        document.add_heading("Contents", level=1)
        with self.assertRaisesRegex(MODULE.ModelError, "fresh page"):
            MODULE._assert_contents_starts_on_fresh_page(document)

    def test_persisted_layout_policy_rejects_drift(self) -> None:
        policy, digest = MODULE.load_layout_policy()
        self.assertEqual(policy["policy_id"], "authoritative-report-default")
        self.assertRegex(digest, r"^[0-9a-f]{64}$")
        changed = dict(policy)
        changed["cover"] = dict(policy["cover"])
        changed["cover"]["classification_location"] = "header_and_body"
        with self.assertRaisesRegex(MODULE.ModelError, "cover rules"):
            MODULE.validate_layout_policy(changed)

    def test_connector_routes_terminate_at_node_boundaries(self) -> None:
        source = (80, 180, 410, 330)
        target = (920, 480, 1250, 630)
        points = MODULE._route_connector(source, target, 0, 6, 4)
        self.assertGreaterEqual(len(points), 4)
        self.assertEqual(points[0][1], source[3])
        self.assertEqual(points[-1][1], target[1])
        self.assertTrue(all(not (source[0] < x < source[2] and source[1] < y < source[3]) for x, y in points[1:]))
        self.assertTrue(all(not (target[0] < x < target[2] and target[1] < y < target[3]) for x, y in points[:-1]))

    def test_model_rejects_prohibited_internal_lessons_content(self) -> None:
        model = _model()
        model["review_trigger"] = "See LL-0001"
        with self.assertRaisesRegex(MODULE.ModelError, "prohibited"):
            MODULE.validate_report_model(model)

    def test_model_rejects_private_url_and_duplicate_reference(self) -> None:
        model = _model()
        model["references"][0]["url"] = "https://127.0.0.1/private"
        with self.assertRaisesRegex(MODULE.ModelError, "non-public"):
            MODULE.validate_report_model(model)
        model = _model()
        model["references"][1]["id"] = "REF-001"
        with self.assertRaisesRegex(MODULE.ModelError, "duplicate reference"):
            MODULE.validate_report_model(model)

    def test_model_rejects_ipv6_loopback_url(self) -> None:
        model = _model()
        model["references"][0]["url"] = "https://[::1]/private"
        with self.assertRaisesRegex(MODULE.ModelError, "non-public"):
            MODULE.validate_report_model(model)

    def test_model_rejects_legacy_loopback_url_spellings(self) -> None:
        for hostname in ("2130706433", "0x7f000001", "127.1"):
            with self.subTest(hostname=hostname):
                model = _model()
                model["references"][0]["url"] = f"https://{hostname}/private"
                with self.assertRaisesRegex(MODULE.ModelError, "non-public"):
                    MODULE.validate_report_model(model)

    def test_builds_candidate_that_passes_authoritative_validator(self) -> None:
        from docx import Document
        from validate_assessment_report import validate_report

        model = MODULE.validate_report_model(_model())
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            output = root / "assessment.docx"
            figure = root / "figure.png"
            MODULE.build_report(model, output, figure)
            result = validate_report(output)
            rendered = Document(output)
            contents_index = next(index for index, paragraph in enumerate(rendered.paragraphs) if paragraph.text == "Contents")
            self.assertIn('w:type="page"', rendered.paragraphs[contents_index - 1]._p.xml)
            header_text = " ".join(paragraph.text for paragraph in rendered.sections[0].header.paragraphs)
            self.assertIn(str(model["assessment"]), header_text)
            self.assertNotIn(str(model["extension_id"]), header_text)
            self.assertIn("Classification: PUBLIC", header_text)
            footer_text = "\n".join(
                paragraph.text for section in rendered.sections for paragraph in section.footer.paragraphs
            )
            self.assertIn("Page ", footer_text)
            self.assertNotIn("PUBLIC", footer_text)
            body_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
            self.assertNotIn("Classification: PUBLIC", body_text)
            self.assertNotIn("Classification", [cell.text for table in rendered.tables for row in table.rows for cell in row.cells])
            contents_xml = rendered.paragraphs[contents_index + 1]._p.xml
            self.assertIn('TOC \\o &quot;1-3&quot; \\h \\z \\u', contents_xml)
            self.assertIn('w:dirty="true"', contents_xml)
            self.assertIn("w:updateFields", rendered.settings._element.xml)
            first_heading_after_contents = next(
                index
                for index in range(contents_index + 1, len(rendered.paragraphs))
                if rendered.paragraphs[index].style.name.startswith("Heading")
            )
            self.assertFalse(
                any(
                    paragraph.style.name == "List Bullet"
                    for paragraph in rendered.paragraphs[contents_index + 1:first_heading_after_contents]
                )
            )
            evidence_table = next(
                table
                for table in rendered.tables
                if [cell.text for cell in table.rows[0].cells] == ["ID", "Evidence", "Source", "Method", "State", "Limitation"]
            )
            state_index = 4
            self.assertGreaterEqual(int(evidence_table._tbl.tblGrid.gridCol_lst[state_index].get(MODULE.qn("w:w"))), 1080)
            references_table = next(
                table
                for table in rendered.tables
                if [cell.text for cell in table.rows[0].cells] == ["ID", "Title", "Publisher", "Source", "Accessed", "Applicability"]
            )
            accessed_index = 4
            self.assertGreaterEqual(int(references_table._tbl.tblGrid.gridCol_lst[accessed_index].get(MODULE.qn("w:w"))), 1260)
            self.assertTrue(
                all(row.cells[accessed_index]._tc.tcPr.first_child_found_in("w:noWrap") is not None for row in references_table.rows)
            )
        self.assertTrue(result.passed, result.failures)
        assert result.metrics
        self.assertGreaterEqual(result.metrics.words, 4_000)
        self.assertGreaterEqual(result.metrics.headings, 40)
        self.assertGreaterEqual(result.metrics.tables, 15)
        self.assertGreaterEqual(result.metrics.reference_bookmarks, 12)
        self.assertGreaterEqual(result.metrics.external_links, 12)


if __name__ == "__main__":
    unittest.main()
